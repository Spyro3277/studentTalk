from fastapi import FastAPI, WebSocket, WebSocketDisconnect, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse
import ollama
import asyncio
import json
import PyPDF2
from sentence_transformers import SentenceTransformer
import numpy as np
from textblob import TextBlob
from vaderSentiment.vaderSentiment import SentimentIntensityAnalyzer
import datetime
import os
from typing import List, Dict, Any
import faiss
import docx
import pickle
import wellbeingMonitor
import courseKnowledgeBase
from io import BytesIO

#http://0.0.0.0:8000/wellbeing_dashboard

app = FastAPI(title="Student Chatbot")

# CORS middleware for frontend
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Initialize models and analyzers
sentence_model = SentenceTransformer('all-MiniLM-L6-v2')
sentiment_analyzer = SentimentIntensityAnalyzer()

# In-memory storage (replace with database later)
course_content = {}
chat_history = []
student_interactions = {}
wellbeing_flags = []


#Initialize wellbeing monitor and knowledgebase
wellbeing_monitor = wellbeingMonitor.WellbeingMonitor()

knowledge_base = courseKnowledgeBase.CourseKnowledgeBase()


@app.post("/upload_syllabus")
async def upload_syllabus(file: UploadFile = File(...)):
    
    try:
        content = await file.read()
        
        if file.filename.endswith('.pdf'):
            # Extract text from PDF
            pdf_reader = PyPDF2.PdfReader(BytesIO(content))
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text()
        elif file.filename.endswith('.docx'):
            #Extract text from .docx files
            document = docx.Document(BytesIO(content))
            text = ""

            for paragraph in document.paragraphs:
                text += paragraph.text 
        else:
            text = content.decode('utf-8')
        
        knowledge_base.add_document(text, file.filename)
        return {"message": f"Successfully uploaded {file.filename}"}
    
    except Exception as e:
        return {"error": str(e)}
    

@app.post("/uploadAssignment")
async def uploadAssignment(file: UploadFile = File(...)):
    try:
        content = await file.read()
        
        if file.filename.endswith('.pdf'):
            # Extract text from PDF
            pdf_reader = PyPDF2.PdfReader(BytesIO(content))
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text()
        elif file.filename.endswith('.docx'):
            #Extract text from .docx files
            document = docx.Document(BytesIO(content))
            text = ""

            for paragraph in document.paragraphs:
                text += paragraph.text 
        else:
            text = content.decode('utf-8')
        
        knowledge_base.add_document(text, file.filename)
        return {"message": f"Successfully uploaded {file.filename}"}
    
    except Exception as e:
        return {"error": str(e)}
    

@app.get("/wellbeing_dashboard")
async def get_wellbeing_dashboard():
    return {
        "total_flags": len(wellbeing_flags),
        "recent_flags": [flag for flag in wellbeing_flags if 
                        datetime.datetime.fromisoformat(flag['timestamp']) > 
                        datetime.datetime.now() - datetime.timedelta(days=7)],
        "student_summary": {}
    }

@app.websocket("/ws/{student_id}")
async def websocket_endpoint(websocket: WebSocket, student_id: str):
    await websocket.accept()
    
    if student_id not in student_interactions:
        student_interactions[student_id] = []
    
    try:
        while True:
            # Receive message from student
            data = await websocket.receive_text()
            message_data = json.loads(data)
            user_message = message_data.get('message', '')
            
            # Log interaction
            interaction = {
                'timestamp': datetime.datetime.now().isoformat(),
                'student_id': student_id,
                'message': user_message,
                'type': 'user'
            }
            student_interactions[student_id].append(interaction)
            chat_history.append(interaction)
            
            # Analyze wellbeing
            wellbeing_analysis = wellbeing_monitor.analyze_message(user_message, student_id)
            
            # Search knowledge base for relevant content
            relevant_docs = knowledge_base.search_similar(user_message)
            context = "\n".join([doc['content'] for doc in relevant_docs[:2]])
            
            # Generate response using Ollama
            prompt = f"""You are a helpful academic assistant for undergraduate students in C++ programming and algorithms courses.

Context from course materials:
{context}

Student question: {user_message}

Please provide a helpful, encouraging response. If the question is about course content, use the context provided. If it's about programming, provide clear explanations and examples. Always be supportive and understanding of student stress.

Response:"""

            try:
                response = ollama.generate(
                    model='llama3.1:8b',
                    prompt=prompt,
                    stream=False
                )
                bot_response = response['response']
                
            except Exception as e:
                bot_response = "I'm sorry, I'm having trouble processing your request right now. Please try again or contact your instructor."
            
            # Log bot response
            bot_interaction = {
                'timestamp': datetime.datetime.now().isoformat(),
                'student_id': student_id,
                'message': bot_response,
                'type': 'bot',
                'wellbeing_score': wellbeing_analysis['wellbeing_score']
            }
            student_interactions[student_id].append(bot_interaction)
            chat_history.append(bot_interaction)

            
            # Send response back to student
            await websocket.send_text(json.dumps({
                'message': bot_response,
                'wellbeing_score': wellbeing_analysis['wellbeing_score'],
                'flagged': wellbeing_analysis['flag_for_review']
            }))
            
    except WebSocketDisconnect:
        print(f"Student {student_id} disconnected")

# Serve static files (HTML frontend)
app.mount("/static", StaticFiles(directory="static"), name="static")
@app.get("/")
async def get_homepage():
    return FileResponse("static/index.html")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)